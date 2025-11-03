#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertidor de Markdown a Word con formato profesional
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Agregar un hipervínculo a un párrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estilo de hipervínculo
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def setup_styles(doc):
    """Configurar estilos del documento"""
    styles = doc.styles

    # Estilo para código
    try:
        code_style = styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)

    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_font.color.rgb = RGBColor(51, 51, 51)

    code_paragraph = code_style.paragraph_format
    code_paragraph.left_indent = Inches(0.5)
    code_paragraph.space_before = Pt(6)
    code_paragraph.space_after = Pt(6)

    # Estilos de título
    for i in range(1, 7):
        heading = styles[f'Heading {i}']
        heading.font.color.rgb = RGBColor(31, 73, 125)
        if i == 1:
            heading.font.size = Pt(24)
        elif i == 2:
            heading.font.size = Pt(20)
        elif i == 3:
            heading.font.size = Pt(16)
        else:
            heading.font.size = Pt(14)

def parse_markdown_line(line, doc, in_code_block, code_lang, in_table, table_data):
    """Parsear una línea de markdown y agregar al documento"""
    line_stripped = line.rstrip()

    # Bloques de código
    if line_stripped.startswith('```'):
        if in_code_block:
            # Fin del bloque de código
            if code_lang == 'json':
                # Agregar código JSON con formato
                code_text = '\n'.join(table_data)
                p = doc.add_paragraph(code_text, style='Code')
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                # Otro tipo de código
                code_text = '\n'.join(table_data)
                p = doc.add_paragraph(code_text, style='Code')
                p.paragraph_format.left_indent = Inches(0.5)
            return False, '', False, []
        else:
            # Inicio del bloque de código
            lang = line_stripped[3:].strip()
            return True, lang, False, []

    if in_code_block:
        table_data.append(line_stripped)
        return in_code_block, code_lang, in_table, table_data

    # Tablas markdown
    if '|' in line and line.strip().startswith('|'):
        if not in_table:
            # Iniciar nueva tabla
            return in_code_block, code_lang, True, [line_stripped]
        else:
            # Continuar tabla
            table_data.append(line_stripped)
            return in_code_block, code_lang, in_table, table_data
    elif in_table:
        # Finalizar tabla
        if len(table_data) > 2:
            create_table(doc, table_data)
        return in_code_block, code_lang, False, []

    # Títulos
    if line_stripped.startswith('#'):
        level = len(re.match(r'^#+', line_stripped).group())
        text = line_stripped[level:].strip()
        # Remover emojis y limpiar
        text = re.sub(r'[^\w\s\-:¿?¡!().,áéíóúÁÉÍÓÚñÑüÜ]', '', text)
        if text:
            doc.add_heading(text, level=level)
        return in_code_block, code_lang, in_table, table_data

    # Líneas horizontales
    if line_stripped in ['---', '***', '___']:
        doc.add_paragraph('_' * 50)
        return in_code_block, code_lang, in_table, table_data

    # Listas
    if line_stripped.startswith('- ') or line_stripped.startswith('* '):
        text = line_stripped[2:].strip()
        text = clean_markdown_text(text)
        if text:
            p = doc.add_paragraph(text, style='List Bullet')
    elif re.match(r'^\d+\.\s', line_stripped):
        text = re.sub(r'^\d+\.\s', '', line_stripped).strip()
        text = clean_markdown_text(text)
        if text:
            p = doc.add_paragraph(text, style='List Number')
    # Párrafos normales
    elif line_stripped:
        text = clean_markdown_text(line_stripped)
        if text:
            p = doc.add_paragraph(text)

    return in_code_block, code_lang, in_table, table_data

def clean_markdown_text(text):
    """Limpiar texto markdown de formato especial"""
    # Remover marcadores de bold/italic pero mantener el texto
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remover emojis comunes
    text = re.sub(r'[✅❌⚠️🎯📋📘📊💡🔧🐛🎉🔐🚀📖📝📁🌐🔍]', '', text)
    # Limpiar espacios múltiples
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def create_table(doc, table_data):
    """Crear tabla en el documento"""
    if len(table_data) < 2:
        return

    # Parsear filas
    rows = []
    for line in table_data:
        if re.match(r'\|[\s\-:]+\|', line):  # Línea separadora
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Crear tabla
    num_cols = len(rows[0])
    num_rows = len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Llenar tabla
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = clean_markdown_text(cell_text)
                # Encabezados en negrita
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(31, 73, 125)

def convert_md_to_docx(md_file, docx_file):
    """Convertir archivo markdown a Word"""
    print(f"[*] Convirtiendo {md_file} a {docx_file}...")

    # Crear documento
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Configurar estilos
    setup_styles(doc)

    # Leer archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Parsear línea por línea
    in_code_block = False
    code_lang = ''
    in_table = False
    table_data = []

    total_lines = len(lines)
    for idx, line in enumerate(lines):
        if idx % 100 == 0:
            progress = (idx / total_lines) * 100
            print(f"  Progreso: {progress:.1f}% ({idx}/{total_lines} líneas)")

        in_code_block, code_lang, in_table, table_data = parse_markdown_line(
            line, doc, in_code_block, code_lang, in_table, table_data
        )

    # Finalizar tabla si quedó pendiente
    if in_table and len(table_data) > 2:
        create_table(doc, table_data)

    # Finalizar código si quedó pendiente
    if in_code_block and table_data:
        code_text = '\n'.join(table_data)
        doc.add_paragraph(code_text, style='Code')

    # Guardar documento
    doc.save(docx_file)
    print(f"[OK] Documento creado exitosamente: {docx_file}")

if __name__ == '__main__':
    import sys
    import os

    # Archivo de entrada y salida
    md_file = 'DOCUMENTACION_COMPLETA_API_ODOO.md'
    docx_file = 'DOCUMENTACION_COMPLETA_API_ODOO.docx'

    # Verificar que existe el archivo
    if not os.path.exists(md_file):
        print(f"[ERROR] No se encontro el archivo {md_file}")
        sys.exit(1)

    # Convertir
    try:
        convert_md_to_docx(md_file, docx_file)
        print(f"\n[OK] Conversion completada!")
        print(f"[*] Archivo generado: {docx_file}")
    except Exception as e:
        print(f"[ERROR] Error durante la conversion: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

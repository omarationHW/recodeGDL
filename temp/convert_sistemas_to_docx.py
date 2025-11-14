#!/usr/bin/env python3
"""
Script para convertir SISTEMAS_CONFIGURACION_MODULOS_COMPLETO.md a Word (.docx)
Incluye renderizado de diagramas Mermaid a imágenes
"""

import re
import os
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def render_mermaid_to_image(mermaid_code):
    """
    Renderiza código Mermaid a imagen usando mermaid.ink API
    """
    try:
        # Codificar el código Mermaid en base64
        graphbytes = mermaid_code.encode("utf-8")
        base64_bytes = base64.urlsafe_b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")

        # Usar la API de mermaid.ink
        url = f"https://mermaid.ink/img/{base64_string}"

        print(f"   Renderizando diagrama Mermaid... ", end='')
        response = requests.get(url, timeout=30)

        if response.status_code == 200:
            print("OK")
            return BytesIO(response.content)
        else:
            print(f"ERROR (Status: {response.status_code})")
            return None
    except Exception as e:
        print(f"ERROR: {str(e)}")
        return None

def extract_mermaid_blocks(content):
    """
    Extrae todos los bloques de código Mermaid del Markdown
    """
    pattern = r'```mermaid\n(.*?)\n```'
    matches = re.findall(pattern, content, re.DOTALL)
    return matches

def add_heading(doc, text, level=1):
    """
    Agrega un encabezado con formato
    """
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """
    Agrega un párrafo con formato
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_table_from_markdown(doc, table_text):
    """
    Convierte una tabla Markdown a tabla Word
    """
    lines = [line.strip() for line in table_text.split('\n') if line.strip()]

    if len(lines) < 2:
        return

    # Separar header y rows
    header_line = lines[0]
    separator_line = lines[1]
    data_lines = lines[2:] if len(lines) > 2 else []

    # Parsear header
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]

    # Parsear data
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    # Crear tabla en Word
    if not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Agregar headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = header
            # Formato de header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    # Agregar datos
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data

def process_markdown_line(doc, line, in_code_block, code_lines, in_table, table_lines):
    """
    Procesa una línea de Markdown
    """
    # Headers
    if line.startswith('# ') and not in_code_block:
        add_heading(doc, line[2:], level=1)
    elif line.startswith('## ') and not in_code_block:
        add_heading(doc, line[3:], level=2)
    elif line.startswith('### ') and not in_code_block:
        add_heading(doc, line[4:], level=3)
    elif line.startswith('#### ') and not in_code_block:
        add_heading(doc, line[5:], level=4)

    # Separador horizontal
    elif line.strip() == '---' and not in_code_block:
        doc.add_paragraph('_' * 50)

    # Listas con viñetas
    elif line.startswith('- ') and not in_code_block:
        p = doc.add_paragraph(line[2:], style='List Bullet')

    # Texto bold/italic
    elif '**' in line or '*' in line and not in_code_block:
        # Procesar formato inline
        p = doc.add_paragraph()
        # Simplificado: agregar como texto normal
        text = line.replace('**', '').replace('*', '')
        p.add_run(text)

    # Párrafo normal
    elif line.strip() and not in_code_block and not in_table:
        # Limpiar markdown básico
        clean_text = line.replace('**', '').replace('`', '')
        doc.add_paragraph(clean_text)

def convert_md_to_docx(md_file, output_file):
    """
    Convierte el archivo Markdown a Word con diagramas renderizados
    """
    print(f"Leyendo archivo: {md_file}")

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    print("Extrayendo bloques Mermaid...")
    mermaid_blocks = extract_mermaid_blocks(content)
    print(f"   Encontrados: {len(mermaid_blocks)} diagramas\n")

    # Renderizar diagramas Mermaid
    mermaid_images = []
    for i, mermaid_code in enumerate(mermaid_blocks, 1):
        print(f"Procesando diagrama {i}/{len(mermaid_blocks)}")
        image_stream = render_mermaid_to_image(mermaid_code)
        mermaid_images.append(image_stream)

    print("\nCreando documento Word...\n")

    # Crear documento Word
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Procesar línea por línea
    lines = content.split('\n')
    in_code_block = False
    code_type = None
    code_lines = []
    in_table = False
    table_lines = []
    mermaid_idx = 0

    for line in lines:
        # Detectar inicio de bloque de código
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_type = line[3:].strip() or 'text'
                code_lines = []
            else:
                in_code_block = False

                # Si es Mermaid, insertar imagen
                if code_type == 'mermaid':
                    if mermaid_idx < len(mermaid_images) and mermaid_images[mermaid_idx]:
                        try:
                            doc.add_paragraph()  # Espacio antes
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(mermaid_images[mermaid_idx], width=Inches(6))
                            doc.add_paragraph()  # Espacio después
                        except Exception as e:
                            doc.add_paragraph(f"[Error al insertar diagrama: {str(e)}]")
                    else:
                        doc.add_paragraph("[Diagrama Mermaid - no se pudo renderizar]")
                    mermaid_idx += 1

                # Si es SQL u otro código, agregar como código
                elif code_lines:
                    p = doc.add_paragraph()
                    code_text = '\n'.join(code_lines)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    # Fondo gris claro simulado con borde
                    p.paragraph_format.left_indent = Inches(0.5)

                code_lines = []
                code_type = None
            continue

        # Si estamos en bloque de código, acumular líneas
        if in_code_block:
            code_lines.append(line)
            continue

        # Detectar tablas
        if '|' in line and not in_code_block:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        else:
            if in_table and table_lines:
                # Procesar tabla acumulada
                table_text = '\n'.join(table_lines)
                add_table_from_markdown(doc, table_text)
                in_table = False
                table_lines = []

        # Procesar línea normal
        process_markdown_line(doc, line, in_code_block, code_lines, in_table, table_lines)

    # Guardar documento
    print(f"\nGuardando documento: {output_file}")
    doc.save(output_file)
    print("OK Documento Word generado exitosamente\n")

if __name__ == '__main__':
    # Rutas
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)

    # MODIFICADO: Convertir el nuevo documento
    md_file = os.path.join(project_root, 'SISTEMAS_CONFIGURACION_MODULOS_COMPLETO.md')
    output_file = os.path.join(project_root, 'SISTEMAS_CONFIGURACION_MODULOS_COMPLETO.docx')

    print("="*60)
    print("   CONVERSION DE MARKDOWN A WORD")
    print("   Sistemas y Configuracion de Modulos")
    print("="*60 + "\n")

    if not os.path.exists(md_file):
        print(f"ERROR: No se encuentra el archivo {md_file}")
        exit(1)

    convert_md_to_docx(md_file, output_file)

    print("="*60)
    print(f"OK COMPLETADO")
    print(f"Archivo generado: {output_file}")
    print("="*60)
